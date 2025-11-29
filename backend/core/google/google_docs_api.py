import httpx
import tempfile
import io
import re
from pathlib import Path
from typing import Optional, Literal
from html.parser import HTMLParser

from fastapi import APIRouter, HTTPException, Depends, Query
from fastapi.responses import StreamingResponse, Response
from pydantic import BaseModel, Field
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

from core.utils.auth_utils import verify_and_get_user_id_from_jwt
from core.utils.logger import logger
from core.services.supabase import DBConnection
from .google_docs_service import GoogleDocsService
from .google_slides_service import OAuthTokenService

class ConvertToDocsRequest(BaseModel):
    doc_path: str = Field(..., description="Path to the document file in sandbox")
    sandbox_url: str = Field(..., description="URL of the sandbox service")


class ExportDocxRequest(BaseModel):
    content: str = Field(..., description="HTML content to convert to DOCX")
    fileName: str = Field(..., description="Name for the exported file")


class ConvertToDocsResponse(BaseModel):
    success: bool = Field(..., description="Whether the conversion was successful")
    message: str = Field(..., description="Status message")
    docx_url: Optional[str] = Field(None, description="URL to the generated DOCX file (if stored locally)")
    google_docs_url: Optional[str] = Field(None, description="Google Docs URL for the document")
    google_docs_file_id: Optional[str] = Field(None, description="Google Drive file ID")
    is_api_enabled: bool = Field(default=True, description="Whether Google API is enabled")


async def get_db_connection() -> DBConnection:
    db = DBConnection()
    await db.initialize()
    return db


async def get_oauth_service(db: DBConnection = Depends(get_db_connection)) -> OAuthTokenService:
    return OAuthTokenService(db)


async def get_google_docs_service(oauth_service: OAuthTokenService = Depends(get_oauth_service)) -> GoogleDocsService:
    return GoogleDocsService(oauth_service)


main_router = APIRouter()

docs_router = APIRouter(prefix="/document-tools", tags=["document-tools"])

@docs_router.get("/debug")
async def debug_endpoint():
    logger.info("DEBUG: Document tools debug endpoint hit")
    return {"status": "ok", "message": "Document tools router is working"}

@docs_router.post("/convert-and-upload-to-docs", response_model=ConvertToDocsResponse)
async def convert_and_upload_to_google_docs(
    request: ConvertToDocsRequest,
    user_id: str = Depends(verify_and_get_user_id_from_jwt),
    google_service: GoogleDocsService = Depends(get_google_docs_service)
):
    try:
        is_authenticated = await google_service.is_user_authenticated(user_id)

        if not is_authenticated:
            logger.info("User not authenticated with Google, returning auth required response")
            response = ConvertToDocsResponse(
                success=False,
                message="User not authenticated with Google. Please authenticate first.",
                docx_url=None,
                is_api_enabled=False
            )
            return response

        convert_url = f"{request.sandbox_url}/document/convert-to-docx"
        convert_payload = {
            "doc_path": request.doc_path,
            "download": True, 
        }
        logger.info(f"Calling sandbox conversion endpoint: POST {convert_url}")
        logger.debug(f"Conversion payload: {convert_payload}")
        
        async with httpx.AsyncClient(timeout=120.0) as client:
            convert_response = await client.post(
                convert_url,
                json=convert_payload,
                headers={
                    "X-Daytona-Skip-Preview-Warning": "true"
                }
            )
            
            logger.debug(f"Sandbox response status: {convert_response.status_code}")
            
            if not convert_response.is_success:
                try:
                    error_detail = convert_response.json().get("detail", "Unknown error")
                except:
                    error_detail = convert_response.text
                logger.error(f"Sandbox conversion failed: {error_detail}")
                raise HTTPException(
                    status_code=convert_response.status_code,
                    detail=f"DOCX conversion failed: {error_detail}"
                )
            
            docx_content = convert_response.content
            
            filename = "document.docx" 
            content_disposition = convert_response.headers.get("Content-Disposition", "")
            if "filename=" in content_disposition:
                filename = content_disposition.split('filename="')[1].split('"')[0]
        
        logger.info(f"DOCX conversion successful: {filename}")
        
       
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as temp_file:
            temp_docx_path = Path(temp_file.name)
            temp_file.write(docx_content)
            temp_file.flush()  
        
        
        try:
            document_name = filename.replace(".docx", "")
            
            upload_result = await google_service.upload_docx_to_docs(
                docx_file_path=temp_docx_path,
                user_id=user_id,
                document_name=document_name
            )
            return ConvertToDocsResponse(
                success=True,
                message=f"Successfully converted and uploaded to Google Docs",
                docx_url=None, 
                google_docs_url=upload_result["web_view_link"],
                google_docs_file_id=upload_result["file_id"]
            )
            
        finally:
            try:
                temp_docx_path.unlink()
            except:
                pass  
    
    except HTTPException as he:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@docs_router.get("/download")
async def download_document(
    sandbox_url: str = Query(..., description="URL of the sandbox service"),
    doc_path: str = Query(..., description="Path to the document file in sandbox"),
    format: Literal["pdf", "docx"] = Query(..., description="Download format: pdf or docx"),
    document_name: str = Query(default="document", description="Name for the downloaded file"),
    _user_id: str = Depends(verify_and_get_user_id_from_jwt)
):
    """
    Proxy endpoint for document downloads.
    This bypasses browser CORS restrictions by making the request from the backend.
    """
    try:
        logger.info(f"Document download request: format={format}, doc_path={doc_path}")
        
        convert_url = f"{sandbox_url}/document/convert-to-{format}"
        
        async with httpx.AsyncClient(timeout=120.0) as client:
            response = await client.post(
                convert_url,
                json={
                    "doc_path": doc_path,
                    "download": True
                },
                headers={
                    "X-Daytona-Skip-Preview-Warning": "true"
                }
            )
            
            if not response.is_success:
                try:
                    error_detail = response.json().get("detail", "Unknown error")
                except:
                    error_detail = response.text[:200] if response.text else "Unknown error"
                logger.error(f"Document conversion failed: {error_detail}")
                raise HTTPException(
                    status_code=response.status_code,
                    detail=f"Document conversion failed: {error_detail}"
                )
            
            content = response.content
            content_type = response.headers.get("content-type", "application/octet-stream")
            
            # Determine proper content type and extension
            if format == "pdf":
                content_type = "application/pdf"
                filename = f"{document_name}.pdf"
            else:
                content_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                filename = f"{document_name}.docx"
            
            logger.info(f"Successfully converted document to {format}, size: {len(content)} bytes")
            
            return StreamingResponse(
                iter([content]),
                media_type=content_type,
                headers={
                    "Content-Disposition": f'attachment; filename="{filename}"',
                    "Content-Length": str(len(content))
                }
            )
            
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Document download error: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))


class HTMLToDocxParser(HTMLParser):
    """Simple HTML to DOCX converter using python-docx"""
    
    def __init__(self, document: Document):
        super().__init__()
        self.document = document
        self.current_paragraph = None
        self.current_run = None
        self.list_level = 0
        self.in_list = False
        self.is_bold = False
        self.is_italic = False
        self.is_underline = False
        self.heading_level = 0
        self.in_code = False
        self.in_blockquote = False
        self.text_buffer = ""
        
    def handle_starttag(self, tag, attrs):
        tag = tag.lower()
        
        if tag in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
            self._flush_text()
            self.heading_level = int(tag[1])
            self.current_paragraph = self.document.add_heading(level=self.heading_level)
            
        elif tag == 'p':
            self._flush_text()
            self.current_paragraph = self.document.add_paragraph()
            if self.in_blockquote:
                self.current_paragraph.paragraph_format.left_indent = Inches(0.5)
                
        elif tag in ['strong', 'b']:
            self._flush_text()
            self.is_bold = True
            
        elif tag in ['em', 'i']:
            self._flush_text()
            self.is_italic = True
            
        elif tag == 'u':
            self._flush_text()
            self.is_underline = True
            
        elif tag == 'br':
            if self.current_paragraph:
                self.current_paragraph.add_run('\n')
                
        elif tag in ['ul', 'ol']:
            self._flush_text()
            self.in_list = True
            self.list_level += 1
            
        elif tag == 'li':
            self._flush_text()
            self.current_paragraph = self.document.add_paragraph(style='List Bullet')
            
        elif tag == 'blockquote':
            self._flush_text()
            self.in_blockquote = True
            
        elif tag in ['code', 'pre']:
            self._flush_text()
            self.in_code = True
            if tag == 'pre':
                self.current_paragraph = self.document.add_paragraph()
                
        elif tag == 'hr':
            self._flush_text()
            self.document.add_paragraph('_' * 50)
            
        elif tag == 'a':
            # Just continue with text, links become plain text in DOCX
            pass
            
    def handle_endtag(self, tag):
        tag = tag.lower()
        
        if tag in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
            self._flush_text()
            self.heading_level = 0
            self.current_paragraph = None
            
        elif tag == 'p':
            self._flush_text()
            self.current_paragraph = None
            
        elif tag in ['strong', 'b']:
            self._flush_text()
            self.is_bold = False
            
        elif tag in ['em', 'i']:
            self._flush_text()
            self.is_italic = False
            
        elif tag == 'u':
            self._flush_text()
            self.is_underline = False
            
        elif tag in ['ul', 'ol']:
            self._flush_text()
            self.list_level -= 1
            if self.list_level <= 0:
                self.in_list = False
                self.list_level = 0
                
        elif tag == 'li':
            self._flush_text()
            self.current_paragraph = None
            
        elif tag == 'blockquote':
            self._flush_text()
            self.in_blockquote = False
            
        elif tag in ['code', 'pre']:
            self._flush_text()
            self.in_code = False
            if tag == 'pre':
                self.current_paragraph = None
                
    def handle_data(self, data):
        # Clean up whitespace but preserve meaningful content
        text = data
        if not self.in_code:
            text = re.sub(r'\s+', ' ', text)
        
        if text.strip() or (self.in_code and text):
            self.text_buffer += text
            
    def _flush_text(self):
        if not self.text_buffer:
            return
            
        text = self.text_buffer
        self.text_buffer = ""
        
        if not self.current_paragraph:
            self.current_paragraph = self.document.add_paragraph()
            
        run = self.current_paragraph.add_run(text)
        run.bold = self.is_bold
        run.italic = self.is_italic
        run.underline = self.is_underline
        
        if self.in_code:
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
            
    def finalize(self):
        self._flush_text()


def html_to_docx(html_content: str, title: str = "Document") -> bytes:
    """Convert HTML content to DOCX bytes"""
    document = Document()
    
    # Set document properties
    core_props = document.core_properties
    core_props.title = title
    core_props.author = "Suna AI"
    
    # Parse HTML and build document
    parser = HTMLToDocxParser(document)
    parser.feed(html_content)
    parser.finalize()
    
    # Save to bytes
    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


# Export router for the /api/export/docx endpoint (matches Next.js route)
export_router = APIRouter(prefix="/export", tags=["export"])


@export_router.post("/docx")
async def export_html_to_docx(request: ExportDocxRequest):
    """
    Convert HTML content to DOCX file.
    This endpoint mirrors the Next.js /api/export/docx route.
    """
    try:
        logger.info(f"DOCX export request for: {request.fileName}")
        
        if not request.content or not request.fileName:
            raise HTTPException(
                status_code=400,
                detail="Content and fileName are required"
            )
        
        # Convert HTML to DOCX
        docx_bytes = html_to_docx(request.content, request.fileName)
        
        logger.info(f"Successfully generated DOCX: {request.fileName}.docx ({len(docx_bytes)} bytes)")
        
        return Response(
            content=docx_bytes,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={
                "Content-Disposition": f'attachment; filename="{request.fileName}.docx"',
                "Content-Length": str(len(docx_bytes))
            }
        )
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"DOCX export error: {str(e)}")
        raise HTTPException(
            status_code=500,
            detail=f"Failed to generate DOCX file: {str(e)}"
        )


main_router.include_router(docs_router)
main_router.include_router(export_router)
router = main_router
